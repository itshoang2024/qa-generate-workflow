from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable
from uuid import uuid4

from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pydantic import BaseModel

from app.domain.models import GDDSection


class ParsedGDD(BaseModel):
    sections: list[GDDSection]


def parse_docx_gdd(path: Path, run_id: str) -> ParsedGDD:
    from docx import Document

    if not path.exists():
        raise FileNotFoundError(f"GDD file not found: {path}")

    document = Document(str(path))
    sections: list[GDDSection] = []
    current: GDDSection | None = None
    parent_stack: dict[int, str] = {}

    for block in _iter_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style = block.style.name if block.style else ""
            if style.startswith("Heading"):
                level = _heading_level(style)
                section_id, title = _parse_heading(text)
                parent_id = _parent_for_level(parent_stack, level)
                current = GDDSection(
                    id=f"sec_{uuid4().hex[:12]}",
                    run_id=run_id,
                    section_id=section_id,
                    title=title,
                    level=level,
                    parent_id=parent_id,
                )
                sections.append(current)
                parent_stack[level] = section_id
                for deeper_level in [key for key in parent_stack if key > level]:
                    parent_stack.pop(deeper_level, None)
                continue

            if current is None:
                continue
            current.text = _join_text(current.text, text)
            if _is_special_note(text):
                current.flags.append(text)

        if isinstance(block, Table) and current is not None:
            current.tables.append(_table_to_rows(block))

    for section in sections:
        actionable, reason = classify_actionability(section)
        section.actionable = actionable
        section.actionability_reason = reason

    _reclassify_containers(sections)

    return ParsedGDD(sections=sections)


def _reclassify_containers(sections: list[GDDSection]) -> None:
    """Second-pass: sections marked insufficient_text that own actionable children
    are structural containers, not thin sections.  Relabel so operators can
    distinguish "truly empty" from "parent heading with rich subsections".
    Actionability (False) is preserved — containers have no direct prose to send
    to agents; their children are already included in the actionable set.
    """
    actionable_parent_ids: set[str] = {
        s.parent_id
        for s in sections
        if s.actionable and s.parent_id is not None
    }
    for section in sections:
        if (
            not section.actionable
            and section.actionability_reason == "insufficient_text"
            and section.section_id in actionable_parent_ids
        ):
            section.actionability_reason = "container"


def classify_actionability(section: GDDSection) -> tuple[bool, str | None]:
    text = section.text.lower()
    if section.section_id in {"§1", "§15"}:
        return False, "metadata"
    if section.section_id == "§13" or "project-context.md" in text:
        return False, "external_dependency"
    if len(text.split()) < 8 and not section.tables:
        return False, "insufficient_text"
    return True, None


def _iter_blocks(document: DocxDocument) -> Iterable[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _heading_level(style_name: str) -> int:
    match = re.search(r"Heading\s+(\d+)", style_name)
    return int(match.group(1)) if match else 1


def _parse_heading(text: str) -> tuple[str, str]:
    match = re.match(r"^(\d+(?:\.\d+)*)(?:\.)?\s+(.+)$", text)
    if not match:
        return f"§{text[:20].strip().replace(' ', '-').lower()}", text
    return f"§{match.group(1)}", match.group(2).strip()


def _parent_for_level(parent_stack: dict[int, str], level: int) -> str | None:
    for candidate_level in range(level - 1, 0, -1):
        if candidate_level in parent_stack:
            return parent_stack[candidate_level]
    return None


def _join_text(existing: str, new_text: str) -> str:
    return new_text if not existing else f"{existing}\n{new_text}"


def _is_special_note(text: str) -> bool:
    return text.startswith(("IMPORTANT:", "NOTE:", "TIP:"))


def _table_to_rows(table: Table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]
